'''
Author : Aditya Bhatt 24-02-2025 

NOTE:
1.We are to create a deep research clone

To enhance our deep research capabilities and align more closely with OpenAI's offerings, the following improvements are recommended:
1. Implement a more advanced web scraper to gather data effectively.
2. Integrate a specialized LLM for report generation, potentially a fine-tuned model designed specifically for crafting detailed reports.
3. While the current system excels at summarization, it lacks the depth required for comprehensive analysis.

OpenAI's fine-tuned model (o3) for deep research tasks, particularly in browser-based research and report creation, represents a significant competitive advantage. However, it is essential to recognize that the internet is an open resource, and we can also strive to achieve a similar level of excellence.

This is not to undermine OpenAI's achievements; their deep research capabilities are indeed remarkable. Nonetheless, we believe that with the right enhancements, we can approach that standard.

TODO:
1.New approach to research.
(query->user_query->research_topics->scrape_urls->scrape_content->(each section will go in a llm call)get two things headers and content)

(Once we have the headers and content we can store them as a variable)

(In the last generate a report using the headers and content)

'''

from http import client
from client.azure import Client
from openai import AzureOpenAI
from schemas.user_query import UserQuery, ResearchQuery
from client.serp import get_search_links
from utils.link_scrapper import extract_content_from_urls
import json
import os
import re
import time
from datetime import datetime
from utils.token_tracker import initialize_trackers, update_and_display_metrics
from utils.user_query import get_search_query
from utils.research_query import get_research_queries
from utils.summary import summarize_content, clean_content
from utils.json_saver import save_summaries_to_json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

input_token_tracker, output_token_tracker = initialize_trackers()

def main():
    global input_token_tracker
    global output_token_tracker
    
    start_time = time.time()
    print(f"Script started at: {datetime.fromtimestamp(start_time).strftime('%Y-%m-%d %H:%M:%S')}")
    
    input_query = input("Enter the topic on which you want to do research: ")
    print("User Wants to research on: ", input_query)

    print("Getting Search Query based on user query...")
    get_client=Client()
    client=get_client.get_azure_client()

    print("Azure Client Created Successfully,Azure OpenAI will be used for this code")

    search_query, input_token_tracker, output_token_tracker = get_search_query(client, input_query, input_token_tracker, output_token_tracker)

    print("Search Query Generated Successfully")

    print("Search Query: ", search_query.query)
    current_cost_inr = update_and_display_metrics(input_token_tracker, output_token_tracker, "mini")

    print("Input Token Usage: ", input_token_tracker)
    print("Output Token Usage: ", output_token_tracker)
    print(f"Current Cost in ₹: {current_cost_inr:.5f}")


    print("The process of generating research queries is starting...")

    # Based on the search query,break it into subtopics and generate research queries
    research_queries, input_token_tracker, output_token_tracker = get_research_queries(client, search_query.query, input_token_tracker, output_token_tracker)

    print("Research Queries Generated Successfully")

    print("Research Queries: ", research_queries.queries)

    current_cost_inr = update_and_display_metrics(input_token_tracker, output_token_tracker, "mini")

    print("Input Token Usage: ", input_token_tracker)
    print("Output Token Usage: ", output_token_tracker)
    print(f"Current Cost in ₹: {current_cost_inr:.5f}")
    research_queries_list=research_queries.queries[0:2]
    # print("For this run we are taking only 2 research queries to save cost and response time")
    # print("Research Queries: ", research_queries_list)
    master_data=[]
    for query in research_queries_list:

        query_content=''
        print("Research Query: ", query)

        print("Processing the research query...")

         #Get the search links based on the research query
        search_links = get_search_links(query)

        print("Search Links Generated Successfully")

        print("Search Links: ", search_links)

        #Get the content from the search links
        raw_data_dict = extract_content_from_urls(search_links)

        print("Links Scraped Successfully and raw data is stored in a dictionary")

        #Pass the raw data to the LLM to get the headers and content,in this dict we have url as key and content as value

        for url,content in raw_data_dict.items():
            print("URL: ", url)
            content_exactarcted ,input_token_tracker, output_token_tracker = summarize_content(client,url,content,input_token_tracker, output_token_tracker)

            print("Content Extracted : ", content_exactarcted[0:100])


            current_cost_inr = update_and_display_metrics(input_token_tracker, output_token_tracker, "mini")
            print("Input Token Usage: ", input_token_tracker)
            print("Output Token Usage: ", output_token_tracker)
            print(f"Current Cost in ₹: {current_cost_inr:.5f}")
            print("Cleaning the content...")
            cleaned_content, input_token_tracker, output_token_tracker = clean_content(client, content_exactarcted, input_token_tracker, output_token_tracker,query)
            query_content+=cleaned_content
            print("Content Cleaned Successfully")
            print("Cleaned Content: ", cleaned_content)

        


        master_data.append({
            "header": query,
            "content": query_content,
            "url": search_links
        })

    print("Master Data Generated Successfully")
    with open("master_data.json", "w") as f:
        json.dump(master_data, f)
    
    # Generate Word document from master_data
    create_word_document(master_data, input_query)
    
    end_time = time.time()
    execution_time = end_time - start_time
    print(f"Script ended at: {datetime.fromtimestamp(end_time).strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Total execution time: {execution_time:.2f} seconds ({execution_time/60:.2f} minutes)")
       
    return master_data

def create_word_document(master_data, research_topic):
    """
    Create a Word document from the master data
    """
    print("Creating Word document from master data...")
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f"Research Report: {research_topic}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.italic = True
    
    # Add a page break after title page
    doc.add_page_break()
    
    # Add table of contents header
    doc.add_heading("Table of Contents", level=1)
    
    # Simple table of contents (just the section headers)
    for i, section in enumerate(master_data, 1):
        toc_para = doc.add_paragraph()
        toc_para.add_run(f"{i}. {section['header']}").bold = True
    
    # Add a page break after TOC
    doc.add_page_break()
    
    # Add content sections
    for section in master_data:
        # Add section header
        doc.add_heading(section['header'], level=1)
        
        # Process content to identify and format subsections
        content = section['content']
        
        # Remove markdown formatting like ** from the content
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
        
        # Split content by potential subsection headers (lines ending with colon)
        parts = re.split(r'([\w\s]+:)\n', content)
        
        if len(parts) > 1:
            # First part might be introductory text
            if parts[0].strip():
                doc.add_paragraph(parts[0].strip())
            
            # Process subsections
            for i in range(1, len(parts), 2):
                if i+1 < len(parts):
                    # Add subsection header
                    subsection_title = parts[i].strip()
                    doc.add_heading(subsection_title, level=2)
                    
                    # Add subsection content
                    subsection_content = parts[i+1].strip()
                    doc.add_paragraph(subsection_content)
        else:
            # If no subsections detected, add the whole content
            doc.add_paragraph(content)
        
        # Add sources
        doc.add_heading("Sources", level=2)
        for url in section['url']:
            doc.add_paragraph(url, style='List Bullet')
        
        # Add page break between sections
        doc.add_page_break()
    
    # Ensure reports directory exists
    os.makedirs("reports", exist_ok=True)
    
    # Save the document
    filename = f"reports/research_report_{research_topic}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    print(f"Word document created successfully: {filename}")

if __name__ == "__main__":
    main()
